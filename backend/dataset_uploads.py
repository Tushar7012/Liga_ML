"""Helpers for session-scoped dataset uploads to the Hugging Face Hub."""

import asyncio
import csv
import io
import json
import os
import re
import uuid
from dataclasses import dataclass
from typing import Any
from urllib.parse import quote

from fastapi import HTTPException, UploadFile
from huggingface_hub import HfApi
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

MAX_DATASET_UPLOAD_BYTES = 100 * 1024 * 1024
ALLOWED_DATASET_EXTENSIONS = {"csv", "json", "jsonl", "pdf", "docx", "xlsx"}
NORMALIZED_DATASET_FILENAME = "train.jsonl"
NORMALIZED_DATASET_FORMAT = "jsonl"
_TRAINING_TOP_LEVEL_KEYS = {
    "messages",
    "prompt",
    "completion",
    "instruction",
    "output",
    "input",
    "response",
    "question",
    "answer",
}
_SAFE_FILENAME_RE = re.compile(r"[^A-Za-z0-9._-]+")
_SAFE_NAMESPACE_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9._-]{0,95}$")


@dataclass(frozen=True)
class DatasetUpload:
    session_id: str
    repo_id: str
    repo_type: str
    private: bool
    upload_id: str
    config_name: str
    filename: str
    original_filename: str
    path_in_repo: str
    raw_path_in_repo: str
    normalized_path_in_repo: str
    normalized_format: str
    normalized_row_count: int
    source_format: str
    supports_training: bool
    size_bytes: int
    format: str
    hub_url: str
    load_dataset_snippet: str

    def response_payload(self) -> dict[str, str | int | bool]:
        return {
            "session_id": self.session_id,
            "repo_id": self.repo_id,
            "repo_type": self.repo_type,
            "private": self.private,
            "upload_id": self.upload_id,
            "config_name": self.config_name,
            "filename": self.filename,
            "path_in_repo": self.path_in_repo,
            "raw_path_in_repo": self.raw_path_in_repo,
            "normalized_path_in_repo": self.normalized_path_in_repo,
            "normalized_format": self.normalized_format,
            "normalized_row_count": self.normalized_row_count,
            "source_format": self.source_format,
            "supports_training": self.supports_training,
            "size_bytes": self.size_bytes,
            "format": self.format,
            "hub_url": self.hub_url,
            "load_dataset_snippet": self.load_dataset_snippet,
        }


@dataclass(frozen=True)
class DatasetRepoCardEntry:
    upload_id: str
    config_name: str
    raw_path: str
    normalized_path: str
    source_format: str
    normalized_row_count: int


def sanitize_dataset_filename(filename: str | None) -> str:
    """Return a Hub-safe basename while preserving the extension."""
    raw = os.path.basename(filename or "").strip()
    if not raw:
        raw = "dataset.csv"

    safe = _SAFE_FILENAME_RE.sub("-", raw).strip(".-_")
    if not safe:
        safe = "dataset.csv"

    stem, ext = os.path.splitext(safe)
    if not stem:
        stem = "dataset"
    if not ext:
        ext = ".csv"

    max_stem_len = 96 - len(ext)
    stem = stem[:max_stem_len].strip(".-_") or "dataset"
    return f"{stem}{ext.lower()}"


def display_filename(filename: str | None, fallback: str) -> str:
    raw = os.path.basename(filename or "").strip()
    if not raw:
        return fallback
    cleaned = "".join(char for char in raw if ord(char) >= 32)
    return cleaned[:160] or fallback


def dataset_format_from_filename(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    if ext not in ALLOWED_DATASET_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=(
                "Only .csv, .json, .jsonl, .pdf, .docx, and .xlsx dataset "
                "files are supported."
            ),
        )
    return ext


def session_dataset_repo_id(hf_username: str | None, session_id: str) -> str:
    namespace = (hf_username or "").strip()
    if not namespace or not _SAFE_NAMESPACE_RE.fullmatch(namespace):
        raise HTTPException(
            status_code=400,
            detail="Could not determine a valid Hugging Face namespace.",
        )

    safe_session_id = re.sub(r"[^A-Za-z0-9]+", "-", session_id).strip("-")
    if not safe_session_id:
        safe_session_id = uuid.uuid4().hex[:8]
    return f"{namespace}/ml-intern-{safe_session_id[:8]}-datasets"


async def upload_size_bytes(upload: UploadFile) -> int:
    await asyncio.to_thread(upload.file.seek, 0, os.SEEK_END)
    size = await asyncio.to_thread(upload.file.tell)
    await asyncio.to_thread(upload.file.seek, 0)
    return int(size)


async def validate_dataset_upload(upload: UploadFile) -> tuple[str, str, int]:
    dataset_format = dataset_format_from_filename(upload.filename or "")
    safe_filename = sanitize_dataset_filename(upload.filename)
    size = await upload_size_bytes(upload)
    if size <= 0:
        raise HTTPException(status_code=400, detail="Uploaded dataset file is empty.")
    if size > MAX_DATASET_UPLOAD_BYTES:
        raise HTTPException(
            status_code=413,
            detail="Dataset upload exceeds the 100 MB limit.",
        )
    return safe_filename, dataset_format, size


def _bad_dataset(detail: str) -> None:
    raise HTTPException(status_code=400, detail=detail)


def _json_safe(value: Any) -> Any:
    if isinstance(value, dict):
        return {str(key): _json_safe(child) for key, child in value.items()}
    if isinstance(value, list):
        return [_json_safe(child) for child in value]
    if isinstance(value, tuple):
        return [_json_safe(child) for child in value]
    if isinstance(value, bytes):
        try:
            return value.decode("utf-8")
        except UnicodeDecodeError:
            return value.hex()
    return value


def _cell_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _chunk_text(text: str, *, min_size: int = 1000, max_size: int = 2000) -> list[str]:
    clean = re.sub(r"\s+", " ", text).strip()
    if not clean:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(clean):
        end = min(start + max_size, len(clean))
        if end < len(clean):
            split = clean.rfind(" ", start + min_size, end)
            if split > start:
                end = split
        chunk = clean[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end
    return chunks


def _messages_text(messages: Any) -> str:
    if not isinstance(messages, list):
        return ""
    parts: list[str] = []
    for message in messages:
        if isinstance(message, dict):
            content = message.get("content")
            if isinstance(content, str) and content.strip():
                parts.append(content.strip())
    return "\n\n".join(parts)


def _row_text_from_json_object(row: dict[str, Any]) -> str:
    messages_text = _messages_text(row.get("messages"))
    if messages_text:
        return messages_text
    for left, right in [
        ("prompt", "completion"),
        ("instruction", "output"),
        ("input", "output"),
        ("input", "response"),
        ("question", "answer"),
    ]:
        left_value = _cell_text(row.get(left))
        right_value = _cell_text(row.get(right))
        if left_value and right_value:
            return f"{left_value}\n\n{right_value}"
    if isinstance(row.get("text"), str) and row["text"].strip():
        return row["text"].strip()
    scalar_parts = [
        f"{key}: {_cell_text(value)}"
        for key, value in row.items()
        if _cell_text(value) and not isinstance(value, (dict, list))
    ]
    return " | ".join(scalar_parts)


def _normalize_json_object(
    row: dict[str, Any],
    *,
    filename: str,
    source_format: str,
    row_index: int,
    line_number: int | None = None,
) -> dict[str, Any]:
    normalized: dict[str, Any] = {
        "source_file": filename,
        "source_format": source_format,
        "row_index": row_index,
        "data": _json_safe(row),
    }
    if line_number is not None:
        normalized["line_number"] = line_number
    for key in _TRAINING_TOP_LEVEL_KEYS:
        if key in row:
            normalized[key] = _json_safe(row[key])
    text = _row_text_from_json_object(row)
    if text:
        normalized["text"] = text
    return normalized


def _normalize_csv(
    contents: bytes, filename: str, source_format: str
) -> list[dict[str, Any]]:
    try:
        text = contents.decode("utf-8-sig")
    except UnicodeDecodeError:
        _bad_dataset(f"{source_format.upper()} dataset must be UTF-8 encoded.")
    reader = csv.DictReader(io.StringIO(text))
    headers = reader.fieldnames or []
    if not headers or any(
        header is None or not str(header).strip() for header in headers
    ):
        _bad_dataset(f"{source_format.upper()} dataset must include a header row.")
    rows: list[dict[str, Any]] = []
    for row in reader:
        cleaned = {str(key).strip(): _cell_text(value) for key, value in row.items()}
        if not any(cleaned.values()):
            continue
        rows.append(
            {
                "text": " | ".join(
                    f"{key}: {value}" for key, value in cleaned.items() if value
                ),
                "source_file": filename,
                "source_format": source_format,
                "row_index": len(rows) + 1,
                "data": cleaned,
            }
        )
    if not rows:
        _bad_dataset(f"{source_format.upper()} dataset contains no usable data rows.")
    return rows


def _json_rows_from_payload(payload: Any) -> list[dict[str, Any]]:
    if isinstance(payload, list):
        rows = payload
    elif isinstance(payload, dict):
        rows = None
        for key in ("data", "rows", "examples"):
            if isinstance(payload.get(key), list):
                rows = payload[key]
                break
        if rows is None:
            rows = [payload]
    else:
        _bad_dataset("JSON dataset must be an object or a list of objects.")

    normalized_rows: list[dict[str, Any]] = []
    for index, row in enumerate(rows, start=1):
        if not isinstance(row, dict):
            _bad_dataset(f"JSON dataset row {index} must be an object.")
        normalized_rows.append(row)
    if not normalized_rows:
        _bad_dataset("JSON dataset contains no rows.")
    return normalized_rows


def _normalize_json(contents: bytes, filename: str) -> list[dict[str, Any]]:
    try:
        payload = json.loads(contents.decode("utf-8-sig"))
    except UnicodeDecodeError:
        _bad_dataset("JSON dataset must be UTF-8 encoded.")
    except json.JSONDecodeError as exc:
        _bad_dataset(f"Invalid JSON dataset: {exc.msg} at line {exc.lineno}.")
    return [
        _normalize_json_object(
            row, filename=filename, source_format="json", row_index=index
        )
        for index, row in enumerate(_json_rows_from_payload(payload), start=1)
    ]


def _normalize_jsonl(contents: bytes, filename: str) -> list[dict[str, Any]]:
    try:
        text = contents.decode("utf-8-sig")
    except UnicodeDecodeError:
        _bad_dataset("JSONL dataset must be UTF-8 encoded.")
    rows: list[dict[str, Any]] = []
    for line_number, line in enumerate(text.splitlines(), start=1):
        if not line.strip():
            continue
        try:
            payload = json.loads(line)
        except json.JSONDecodeError as exc:
            _bad_dataset(f"Invalid JSONL at line {line_number}: {exc.msg}.")
        if not isinstance(payload, dict):
            _bad_dataset(f"JSONL row at line {line_number} must be an object.")
        rows.append(
            _normalize_json_object(
                payload,
                filename=filename,
                source_format="jsonl",
                row_index=len(rows) + 1,
                line_number=line_number,
            )
        )
    if not rows:
        _bad_dataset("JSONL dataset contains no rows.")
    return rows


def _normalize_pdf(contents: bytes, filename: str) -> list[dict[str, Any]]:
    try:
        reader = PdfReader(io.BytesIO(contents))
    except Exception:
        _bad_dataset(
            "PDF dataset has no extractable text. Scanned PDFs are not supported."
        )
    rows: list[dict[str, Any]] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        for chunk in _chunk_text(text):
            rows.append(
                {
                    "text": chunk,
                    "source_file": filename,
                    "source_format": "pdf",
                    "chunk_index": len(rows),
                    "page_number": page_index,
                }
            )
    if not rows:
        _bad_dataset(
            "PDF dataset has no extractable text. Scanned PDFs are not supported."
        )
    return rows


def _normalize_docx(contents: bytes, filename: str) -> list[dict[str, Any]]:
    try:
        document = Document(io.BytesIO(contents))
    except Exception:
        _bad_dataset("DOCX dataset could not be read.")
    text_blocks: list[str] = []
    text_blocks.extend(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    )
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_blocks.append(" | ".join(cells))
    rows: list[dict[str, Any]] = []
    for block in text_blocks:
        for chunk in _chunk_text(block):
            rows.append(
                {
                    "text": chunk,
                    "source_file": filename,
                    "source_format": "docx",
                    "chunk_index": len(rows),
                }
            )
    if not rows:
        _bad_dataset("DOCX dataset contains no extractable text.")
    return rows


def _normalize_xlsx(contents: bytes, filename: str) -> list[dict[str, Any]]:
    try:
        workbook = load_workbook(io.BytesIO(contents), data_only=True, read_only=True)
    except Exception:
        _bad_dataset("XLSX dataset could not be read.")
    rows: list[dict[str, Any]] = []
    for sheet in workbook.worksheets:
        if sheet.sheet_state != "visible":
            continue
        iterator = sheet.iter_rows(values_only=True)
        try:
            header_values = next(iterator)
        except StopIteration:
            continue
        headers = [_cell_text(value) for value in header_values]
        if not headers or not any(headers):
            continue
        for row_number, values in enumerate(iterator, start=2):
            data = {
                header: _cell_text(value)
                for header, value in zip(headers, values, strict=False)
                if header
            }
            if not any(data.values()):
                continue
            rows.append(
                {
                    "text": " | ".join(
                        f"{key}: {value}" for key, value in data.items() if value
                    ),
                    "source_file": filename,
                    "source_format": "xlsx",
                    "row_index": row_number,
                    "sheet_name": sheet.title,
                    "data": data,
                }
            )
    if not rows:
        _bad_dataset("XLSX dataset contains no usable rows.")
    return rows


def normalize_uploaded_dataset(
    contents: bytes,
    filename: str,
    dataset_format: str,
) -> list[dict[str, Any]]:
    if not contents:
        _bad_dataset("Uploaded dataset file is empty.")
    if dataset_format == "csv":
        return _normalize_csv(contents, filename, "csv")
    if dataset_format == "json":
        return _normalize_json(contents, filename)
    if dataset_format == "jsonl":
        return _normalize_jsonl(contents, filename)
    if dataset_format == "pdf":
        return _normalize_pdf(contents, filename)
    if dataset_format == "docx":
        return _normalize_docx(contents, filename)
    if dataset_format == "xlsx":
        return _normalize_xlsx(contents, filename)
    _bad_dataset(f"Unsupported dataset format: {dataset_format}.")


def normalized_rows_to_jsonl(rows: list[dict[str, Any]]) -> bytes:
    if not rows:
        _bad_dataset("Dataset normalization produced no rows.")
    payload = "\n".join(
        json.dumps(_json_safe(row), ensure_ascii=False, sort_keys=True) for row in rows
    )
    return f"{payload}\n".encode("utf-8")


def dataset_hub_url(repo_id: str, path_in_repo: str) -> str:
    quoted_path = quote(path_in_repo, safe="/")
    return f"https://huggingface.co/datasets/{repo_id}/blob/main/{quoted_path}"


def dataset_config_name(upload_id: str) -> str:
    safe_upload_id = re.sub(r"[^A-Za-z0-9]+", "_", upload_id).strip("_").lower()
    if not safe_upload_id:
        safe_upload_id = "dataset"
    return f"upload_{safe_upload_id[:32]}"


def dataset_config_name_from_path(path_in_repo: str) -> str:
    parts = path_in_repo.split("/")
    if len(parts) >= 3 and parts[0] == "uploads":
        return dataset_config_name(parts[1])
    stem = os.path.splitext(os.path.basename(path_in_repo))[0]
    return dataset_config_name(stem)


def is_dataset_upload_path(path_in_repo: str) -> bool:
    parts = path_in_repo.split("/")
    return (
        len(parts) == 3
        and parts[0] == "uploads"
        and bool(parts[1])
        and parts[2] == NORMALIZED_DATASET_FILENAME
    )


def unique_dataset_upload_paths(paths: list[str]) -> list[str]:
    seen = set()
    upload_paths = []
    for path in paths:
        if not is_dataset_upload_path(path) or path in seen:
            continue
        seen.add(path)
        upload_paths.append(path)
    return upload_paths


def dataset_repo_card_entries_from_paths(
    paths: list[str],
) -> list[DatasetRepoCardEntry]:
    by_upload_id: dict[str, dict[str, str]] = {}
    for path in paths:
        parts = path.split("/")
        if len(parts) < 3 or parts[0] != "uploads" or not parts[1]:
            continue
        upload_id = parts[1]
        record = by_upload_id.setdefault(upload_id, {})
        if len(parts) == 3 and parts[2] == NORMALIZED_DATASET_FILENAME:
            record["normalized_path"] = path
        elif len(parts) == 4 and parts[2] == "raw" and parts[3]:
            record["raw_path"] = path

    entries: list[DatasetRepoCardEntry] = []
    for upload_id, record in sorted(by_upload_id.items()):
        normalized_path = record.get("normalized_path")
        if not normalized_path:
            continue
        raw_path = record.get("raw_path", f"uploads/{upload_id}/raw/<original-file>")
        source_format = os.path.splitext(raw_path)[1].lower().lstrip(".") or "unknown"
        entries.append(
            DatasetRepoCardEntry(
                upload_id=upload_id,
                config_name=dataset_config_name(upload_id),
                raw_path=raw_path,
                normalized_path=normalized_path,
                source_format=source_format,
                normalized_row_count=0,
            )
        )
    return entries


def load_dataset_snippet(repo_id: str, config_name: str) -> str:
    return (
        "from datasets import load_dataset\n\n"
        f'dataset = load_dataset("{repo_id}", "{config_name}", '
        'split="train", token=True)'
    )


def dataset_repo_card(
    repo_id: str,
    upload_entries: list[DatasetRepoCardEntry] | list[str],
) -> bytes:
    config_lines = []
    if upload_entries and isinstance(upload_entries[0], str):
        entries = dataset_repo_card_entries_from_paths(upload_entries)  # type: ignore[arg-type]
    else:
        entries = list(upload_entries)  # type: ignore[arg-type]

    unique_entries: list[DatasetRepoCardEntry] = []
    seen_paths: set[str] = set()
    for entry in entries:
        if entry.normalized_path in seen_paths:
            continue
        seen_paths.add(entry.normalized_path)
        unique_entries.append(entry)

    if unique_entries:
        config_lines.append("configs:")
        for entry in unique_entries:
            config_lines.extend(
                [
                    f"- config_name: {entry.config_name}",
                    "  data_files:",
                    "  - split: train",
                    f'    path: "{entry.normalized_path}"',
                ]
            )

    configs = "\n".join(config_lines)
    if configs:
        configs = f"{configs}\n"

    content = f"""---
tags:
- ml-intern
- uploaded-dataset
{configs}---

# {repo_id}

Private dataset files uploaded through ML Intern.

Files are stored under `uploads/<upload_id>/` and are attached to the
corresponding ML Intern session context by Hub reference, not by copying file
contents into the chat.

Each uploaded file is exposed as its own dataset config so files with different
schemas can coexist in the same session repo.

The config paths point to normalized JSONL files for training. Raw source files
are preserved for traceability, including formats that Hugging Face Datasets
does not load directly as train rows.
"""
    if unique_entries:
        content += "\n## Uploaded Files\n\n"
        for entry in unique_entries:
            row_count = (
                f"{entry.normalized_row_count} rows"
                if entry.normalized_row_count
                else "row count unavailable"
            )
            content += f"""### {entry.config_name}

- Raw file: `{entry.raw_path}`
- Normalized train file: `{entry.normalized_path}`
- Original format: `{entry.source_format}`
- Normalized JSONL purpose: train split for `load_dataset(..., "{entry.config_name}")` ({row_count})

"""
    return content.encode("utf-8")


def dataset_context_note(upload: DatasetUpload) -> str:
    return f"""[SYSTEM: The user uploaded a dataset file for this session.

Use this normalized dataset config for HF Jobs or GCP Vertex training when the
task needs the uploaded data. Do not look for the uploaded file on local disk,
do not load the raw file for training, and do not ask the user to upload it
again unless this Hub reference fails.

- Repo ID: {upload.repo_id}
- Repo type: dataset
- Dataset config: {upload.config_name}
- Normalized path: {upload.normalized_path_in_repo}
- Raw path: {upload.raw_path_in_repo}
- Original filename: {upload.original_filename}
- Stored filename: {upload.filename}
- Source format: {upload.source_format}
- Normalized format: {upload.normalized_format}
- Normalized rows: {upload.normalized_row_count}
- Supports training: {upload.supports_training}
- Size: {upload.size_bytes} bytes
- Hub URL: {upload.hub_url}

Load the normalized dataset config with:
```python
{upload.load_dataset_snippet}
```
]"""


async def push_dataset_upload_to_hub(
    *,
    upload: UploadFile,
    session_id: str,
    hf_username: str,
    hf_token: str,
) -> DatasetUpload:
    safe_filename, dataset_format, size = await validate_dataset_upload(upload)
    original_filename = display_filename(upload.filename, safe_filename)
    upload_id = uuid.uuid4().hex[:12]
    config_name = dataset_config_name(upload_id)
    repo_id = session_dataset_repo_id(hf_username, session_id)
    raw_path_in_repo = f"uploads/{upload_id}/raw/{safe_filename}"
    normalized_path_in_repo = f"uploads/{upload_id}/{NORMALIZED_DATASET_FILENAME}"
    path_in_repo = normalized_path_in_repo
    hub_url = dataset_hub_url(repo_id, normalized_path_in_repo)
    snippet = load_dataset_snippet(repo_id, config_name)
    api = HfApi(token=hf_token)

    await asyncio.to_thread(
        api.create_repo,
        repo_id=repo_id,
        repo_type="dataset",
        private=True,
        exist_ok=True,
    )
    await asyncio.to_thread(
        api.update_repo_settings,
        repo_id=repo_id,
        repo_type="dataset",
        private=True,
    )
    await asyncio.to_thread(upload.file.seek, 0)
    file_bytes = await asyncio.to_thread(upload.file.read)
    normalized_rows = normalize_uploaded_dataset(
        file_bytes, safe_filename, dataset_format
    )
    normalized_jsonl = normalized_rows_to_jsonl(normalized_rows)
    repo_files = await asyncio.to_thread(
        api.list_repo_files,
        repo_id=repo_id,
        repo_type="dataset",
    )
    card_entries = dataset_repo_card_entries_from_paths(repo_files)
    card_entries.append(
        DatasetRepoCardEntry(
            upload_id=upload_id,
            config_name=config_name,
            raw_path=raw_path_in_repo,
            normalized_path=normalized_path_in_repo,
            source_format=dataset_format,
            normalized_row_count=len(normalized_rows),
        )
    )
    await asyncio.to_thread(
        api.upload_file,
        path_or_fileobj=file_bytes,
        path_in_repo=raw_path_in_repo,
        repo_id=repo_id,
        repo_type="dataset",
        commit_message=f"Upload raw dataset file {safe_filename}",
    )
    await asyncio.to_thread(
        api.upload_file,
        path_or_fileobj=normalized_jsonl,
        path_in_repo=normalized_path_in_repo,
        repo_id=repo_id,
        repo_type="dataset",
        commit_message=f"Upload normalized dataset rows for {safe_filename}",
    )
    await asyncio.to_thread(
        api.upload_file,
        path_or_fileobj=dataset_repo_card(repo_id, card_entries),
        path_in_repo="README.md",
        repo_id=repo_id,
        repo_type="dataset",
        commit_message="Update ML Intern dataset upload configs",
    )

    return DatasetUpload(
        session_id=session_id,
        repo_id=repo_id,
        repo_type="dataset",
        private=True,
        upload_id=upload_id,
        config_name=config_name,
        filename=safe_filename,
        original_filename=original_filename,
        path_in_repo=path_in_repo,
        raw_path_in_repo=raw_path_in_repo,
        normalized_path_in_repo=normalized_path_in_repo,
        normalized_format=NORMALIZED_DATASET_FORMAT,
        normalized_row_count=len(normalized_rows),
        source_format=dataset_format,
        supports_training=True,
        size_bytes=size,
        format=dataset_format,
        hub_url=hub_url,
        load_dataset_snippet=snippet,
    )

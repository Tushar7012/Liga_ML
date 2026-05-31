import io
import json
import sys
from pathlib import Path
from types import SimpleNamespace

import httpx
import pytest
from docx import Document
from fastapi import HTTPException, UploadFile
from huggingface_hub.errors import HfHubHTTPError
from openpyxl import Workbook
from starlette.datastructures import FormData

_BACKEND_DIR = Path(__file__).resolve().parent.parent.parent / "backend"
if str(_BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(_BACKEND_DIR))

import dataset_uploads  # noqa: E402
from routes import agent  # noqa: E402


def _upload(filename: str, content: bytes = b"a,b\n1,2\n") -> UploadFile:
    return UploadFile(filename=filename, file=io.BytesIO(content))


def _jsonl(rows: list[dict]) -> list[dict]:
    return [
        json.loads(line) for line in rows_to_jsonl(rows).decode("utf-8").splitlines()
    ]


def rows_to_jsonl(rows: list[dict]) -> bytes:
    return b"\n".join(json.dumps(row).encode("utf-8") for row in rows) + b"\n"


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("First paragraph")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Question"
    table.rows[0].cells[1].text = "Answer"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Visible Data"
    sheet.append(["prompt", "completion"])
    sheet.append(["Hi", "Hello"])
    hidden = workbook.create_sheet("Hidden Data")
    hidden.sheet_state = "hidden"
    hidden.append(["prompt", "completion"])
    hidden.append(["ignore", "me"])
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _text_pdf_bytes() -> bytes:
    return b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>
endobj
4 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
5 0 obj
<< /Length 57 >>
stream
BT
/F1 24 Tf
72 720 Td
(Hello from a PDF page) Tj
ET
endstream
endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000241 00000 n
0000000311 00000 n
trailer
<< /Root 1 0 R /Size 6 >>
startxref
417
%%EOF
"""


def _track_close(upload: UploadFile):
    state = {"closed": False}
    original_close = upload.close

    async def close():
        state["closed"] = True
        await original_close()

    upload.close = close
    return state


def _request(
    upload: UploadFile | None = None,
    headers: dict[str, str] | None = None,
):
    state = {"form_called": False}

    class FakeRequest:
        def __init__(self):
            self.headers = headers or {}
            self.cookies = {}

        async def form(self, **_kwargs):
            state["form_called"] = True
            if upload is None:
                raise AssertionError("request.form() should not be called")
            return FormData([("file", upload)])

    return FakeRequest(), state


def test_sanitize_dataset_filename_strips_paths_and_unsafe_chars():
    assert (
        dataset_uploads.sanitize_dataset_filename("../../bad file (final).CSV")
        == "bad-file-final.csv"
    )
    assert dataset_uploads.sanitize_dataset_filename("") == "dataset.csv"


def test_dataset_format_rejects_unsupported_extension():
    for extension in ["csv", "json", "jsonl", "pdf", "docx", "xlsx", "md"]:
        assert (
            dataset_uploads.dataset_format_from_filename(f"rows.{extension}")
            == extension
        )

    with pytest.raises(HTTPException) as exc_info:
        dataset_uploads.dataset_format_from_filename("notes.txt")

    assert exc_info.value.status_code == 400
    assert ".pdf" in exc_info.value.detail

    with pytest.raises(HTTPException):
        dataset_uploads.dataset_format_from_filename("notes")


def test_normalize_csv_rows_to_training_jsonl_schema():
    rows = dataset_uploads.normalize_uploaded_dataset(
        b"\xef\xbb\xbfprompt,completion\nHi,Hello\n",
        "rows.csv",
        "csv",
    )

    assert rows == [
        {
            "text": "prompt: Hi | completion: Hello",
            "source_file": "rows.csv",
            "source_format": "csv",
            "row_index": 1,
            "data": {"prompt": "Hi", "completion": "Hello"},
        }
    ]


def test_normalize_markdown_chunks_to_training_jsonl_schema():
    rows = dataset_uploads.normalize_uploaded_dataset(
        b"# Training Notes\n\nUse concise answers.\n\n## Style\nPrefer examples.",
        "notes.md",
        "md",
    )

    assert rows == [
        {
            "source_format": "md",
            "source_file": "notes.md",
            "chunk_index": 0,
            "text": "# Training Notes Use concise answers. ## Style Prefer examples.",
        }
    ]


def test_normalize_markdown_splits_examples_and_preserves_structured_content():
    rows = dataset_uploads.normalize_uploaded_dataset(
        b"""# Fine-tuning Examples

## Example 1
User: How do I reset my password?
Assistant: Go to Settings, choose Security, then select Reset Password.

## Example 2
- User: How do I export my report?
- Assistant: Open Reports, choose the desired report, and click Export.

| field | value |
| --- | --- |
| source | docs |

```python
print("keep code")
```
""",
        "examples.md",
        "md",
    )

    assert len(rows) == 2
    assert rows[0]["chunk_index"] == 0
    assert rows[0]["text"].startswith("# Fine-tuning Examples\n\n## Example 1")
    assert "User: How do I reset my password?" in rows[0]["text"]
    assert rows[1]["chunk_index"] == 1
    assert rows[1]["text"].startswith("# Fine-tuning Examples\n\n## Example 2")
    assert "- User: How do I export my report?" in rows[1]["text"]
    assert "| source | docs |" in rows[1]["text"]
    assert 'print("keep code")' in rows[1]["text"]


def test_markdown_rows_serialize_as_one_json_object_per_line():
    rows = dataset_uploads.normalize_uploaded_dataset(
        b"# Examples\n\n## One\nUser: A\nAssistant: B\n\n## Two\nUser: C\nAssistant: D",
        "examples.md",
        "md",
    )

    payload = dataset_uploads.normalized_rows_to_jsonl(rows).decode("utf-8")
    lines = payload.splitlines()

    assert len(lines) == 2
    assert [json.loads(line)["chunk_index"] for line in lines] == [0, 1]


def test_normalize_markdown_rejects_empty_text():
    with pytest.raises(HTTPException) as exc_info:
        dataset_uploads.normalize_uploaded_dataset(
            b"   \n\t\n",
            "empty.md",
            "md",
        )

    assert exc_info.value.status_code == 400
    assert "Markdown dataset contains no usable text" in exc_info.value.detail


def test_normalize_json_accepts_lists_named_lists_and_single_objects():
    rows = dataset_uploads.normalize_uploaded_dataset(
        json.dumps(
            {
                "data": [
                    {"messages": [{"role": "user", "content": "hi"}]},
                    {"question": "Q", "answer": "A"},
                ]
            }
        ).encode("utf-8"),
        "rows.json",
        "json",
    )

    assert rows[0]["messages"] == [{"role": "user", "content": "hi"}]
    assert rows[0]["data"] == {"messages": [{"role": "user", "content": "hi"}]}
    assert rows[0]["text"] == "hi"
    assert rows[1]["question"] == "Q"
    assert rows[1]["answer"] == "A"
    assert rows[1]["text"] == "Q\n\nA"

    single = dataset_uploads.normalize_uploaded_dataset(
        b'{"instruction":"Say hi","output":"Hi"}',
        "single.json",
        "json",
    )
    assert single[0]["instruction"] == "Say hi"
    assert single[0]["output"] == "Hi"


def test_normalize_jsonl_reports_invalid_line_number():
    with pytest.raises(HTTPException) as exc_info:
        dataset_uploads.normalize_uploaded_dataset(
            b'{"text":"ok"}\n[]\n',
            "rows.jsonl",
            "jsonl",
        )

    assert exc_info.value.status_code == 400
    assert "line 2" in exc_info.value.detail


def test_normalize_pdf_chunks_extracted_page_text_and_rejects_scanned_pdf():
    rows = dataset_uploads.normalize_uploaded_dataset(
        _text_pdf_bytes(),
        "paper.pdf",
        "pdf",
    )

    assert rows[0]["source_file"] == "paper.pdf"
    assert rows[0]["source_format"] == "pdf"
    assert rows[0]["chunk_index"] == 0
    assert rows[0]["page_number"] == 1
    assert "Hello from a PDF page" in rows[0]["text"]

    with pytest.raises(HTTPException) as exc_info:
        dataset_uploads.normalize_uploaded_dataset(
            b"%PDF-1.4\n%%EOF\n",
            "scan.pdf",
            "pdf",
        )
    assert exc_info.value.status_code == 400
    assert "extractable text" in exc_info.value.detail


def test_normalize_docx_uses_paragraphs_and_table_rows():
    rows = dataset_uploads.normalize_uploaded_dataset(
        _docx_bytes(),
        "notes.docx",
        "docx",
    )

    assert [row["text"] for row in rows] == ["First paragraph", "Question | Answer"]
    assert {row["source_format"] for row in rows} == {"docx"}
    assert [row["chunk_index"] for row in rows] == [0, 1]


def test_normalize_xlsx_uses_visible_sheets_header_and_rows():
    rows = dataset_uploads.normalize_uploaded_dataset(
        _xlsx_bytes(),
        "book.xlsx",
        "xlsx",
    )

    assert rows == [
        {
            "text": "prompt: Hi | completion: Hello",
            "source_file": "book.xlsx",
            "source_format": "xlsx",
            "row_index": 2,
            "sheet_name": "Visible Data",
            "data": {"prompt": "Hi", "completion": "Hello"},
        }
    ]


def test_dataset_repo_card_points_configs_to_normalized_jsonl_and_documents_raw_paths():
    card = dataset_uploads.dataset_repo_card(
        "alice/ml-intern-s1-datasets",
        [
            dataset_uploads.DatasetRepoCardEntry(
                upload_id="oldabc",
                config_name="upload_oldabc",
                raw_path="uploads/oldabc/raw/rows.pdf",
                normalized_path="uploads/oldabc/train.jsonl",
                source_format="pdf",
                normalized_row_count=3,
            ),
            dataset_uploads.DatasetRepoCardEntry(
                upload_id="newdef",
                config_name="upload_newdef",
                raw_path="uploads/newdef/raw/table.csv",
                normalized_path="uploads/newdef/train.jsonl",
                source_format="csv",
                normalized_row_count=2,
            ),
        ],
    ).decode("utf-8")

    assert "configs:" in card
    assert "- config_name: upload_oldabc" in card
    assert '    path: "uploads/oldabc/train.jsonl"' in card
    assert "- config_name: upload_newdef" in card
    assert '    path: "uploads/newdef/train.jsonl"' in card
    assert "Raw file: `uploads/oldabc/raw/rows.pdf`" in card
    assert "Normalized train file: `uploads/oldabc/train.jsonl`" in card
    assert "Original format: `pdf`" in card
    assert card.count("- config_name: upload_oldabc") == 1


@pytest.mark.asyncio
async def test_validate_dataset_upload_rejects_size_over_limit(monkeypatch):
    monkeypatch.setattr(dataset_uploads, "MAX_DATASET_UPLOAD_BYTES", 3)
    upload = _upload("rows.csv", b"abcd")
    try:
        with pytest.raises(HTTPException) as exc_info:
            await dataset_uploads.validate_dataset_upload(upload)
    finally:
        await upload.close()

    assert exc_info.value.status_code == 413


@pytest.mark.asyncio
async def test_push_dataset_upload_creates_private_repo_and_uploads_raw_and_normalized_files(
    monkeypatch,
):
    instances = []

    class FakeApi:
        def __init__(self, token):
            self.token = token
            self.create_calls = []
            self.settings_calls = []
            self.list_calls = []
            self.upload_calls = []
            instances.append(self)

        def create_repo(self, **kwargs):
            self.create_calls.append(kwargs)

        def update_repo_settings(self, **kwargs):
            self.settings_calls.append(kwargs)

        def list_repo_files(self, **kwargs):
            self.list_calls.append(kwargs)
            return [
                "README.md",
                "uploads/oldupload/old.jsonl",
                "uploads/notes.txt",
            ]

        def upload_file(self, **kwargs):
            if kwargs["path_in_repo"].endswith("/raw/Data-Set.csv"):
                assert kwargs["path_or_fileobj"] == b"a,b\n1,2\n"
            if kwargs["path_in_repo"].endswith("/train.jsonl"):
                assert _jsonl(
                    [json.loads(kwargs["path_or_fileobj"].decode("utf-8"))]
                ) == [
                    {
                        "text": "a: 1 | b: 2",
                        "source_file": "Data-Set.csv",
                        "source_format": "csv",
                        "row_index": 1,
                        "data": {"a": "1", "b": "2"},
                    }
                ]
            self.upload_calls.append(kwargs)

    monkeypatch.setattr(dataset_uploads, "HfApi", FakeApi)
    monkeypatch.setattr(
        dataset_uploads.uuid,
        "uuid4",
        lambda: SimpleNamespace(hex="feedfacecafebeef"),
    )

    upload = _upload("../Data Set.CSV")
    try:
        result = await dataset_uploads.push_dataset_upload_to_hub(
            upload=upload,
            session_id="12345678-90ab-cdef-1234-567890abcdef",
            hf_username="alice",
            hf_token="hf-token",
        )
    finally:
        await upload.close()

    api = instances[0]
    assert api.token == "hf-token"
    assert api.create_calls == [
        {
            "repo_id": "alice/ml-intern-12345678-datasets",
            "repo_type": "dataset",
            "private": True,
            "exist_ok": True,
        }
    ]
    assert api.settings_calls == [
        {
            "repo_id": "alice/ml-intern-12345678-datasets",
            "repo_type": "dataset",
            "private": True,
        }
    ]
    assert api.list_calls == [
        {
            "repo_id": "alice/ml-intern-12345678-datasets",
            "repo_type": "dataset",
        }
    ]
    assert [call["path_in_repo"] for call in api.upload_calls] == [
        "uploads/feedfacecafe/raw/Data-Set.csv",
        "uploads/feedfacecafe/train.jsonl",
        "README.md",
    ]
    readme = api.upload_calls[1]["path_or_fileobj"].decode("utf-8")
    readme = api.upload_calls[2]["path_or_fileobj"].decode("utf-8")
    assert "- config_name: upload_feedfacecafe" in readme
    assert '    path: "uploads/feedfacecafe/train.jsonl"' in readme
    assert "Raw file: `uploads/feedfacecafe/raw/Data-Set.csv`" in readme
    assert result.repo_id == "alice/ml-intern-12345678-datasets"
    assert result.config_name == "upload_feedfacecafe"
    assert result.format == "csv"
    assert result.source_format == "csv"
    assert result.raw_path_in_repo == "uploads/feedfacecafe/raw/Data-Set.csv"
    assert result.normalized_path_in_repo == "uploads/feedfacecafe/train.jsonl"
    assert result.path_in_repo == "uploads/feedfacecafe/train.jsonl"
    assert result.normalized_format == "jsonl"
    assert result.normalized_row_count == 1
    assert result.supports_training is True
    assert result.load_dataset_snippet == (
        "from datasets import load_dataset\n\n"
        'dataset = load_dataset("alice/ml-intern-12345678-datasets", '
        '"upload_feedfacecafe", split="train", token=True)'
    )


@pytest.mark.asyncio
async def test_upload_route_requires_hf_token_without_parsing_upload(monkeypatch):
    monkeypatch.delenv("HF_TOKEN", raising=False)
    upload = _upload("rows.csv")
    close_state = _track_close(upload)
    request, request_state = _request(upload)

    async def fake_check_session_access(*_args, **_kwargs):
        return SimpleNamespace(
            is_active=True,
            is_processing=False,
            session=SimpleNamespace(pending_approval=None),
            hf_username="alice",
        )

    monkeypatch.setattr(agent, "_check_session_access", fake_check_session_access)

    try:
        with pytest.raises(HTTPException) as exc_info:
            await agent.upload_session_dataset(
                "s1",
                request,
                {"user_id": "u1", "username": "alice"},
            )

        assert exc_info.value.status_code == 401
        assert request_state["form_called"] is False
        assert close_state["closed"] is False
    finally:
        await upload.close()


@pytest.mark.asyncio
async def test_upload_route_rejects_content_length_before_parsing(monkeypatch):
    upload = _upload("rows.csv")
    close_state = _track_close(upload)
    request, request_state = _request(
        upload,
        headers={
            "content-length": str(
                dataset_uploads.MAX_DATASET_UPLOAD_BYTES
                + agent.DATASET_UPLOAD_MULTIPART_SLACK_BYTES
                + 1
            )
        },
    )

    async def fake_check_session_access(*_args, **_kwargs):
        raise AssertionError("session access should not run for oversized uploads")

    monkeypatch.setattr(agent, "_check_session_access", fake_check_session_access)

    try:
        with pytest.raises(HTTPException) as exc_info:
            await agent.upload_session_dataset(
                "s1",
                request,
                {
                    "user_id": "u1",
                    "username": "alice",
                    agent.INTERNAL_HF_TOKEN_KEY: "hf-token",
                },
            )

        assert exc_info.value.status_code == 413
        assert request_state["form_called"] is False
        assert close_state["closed"] is False
    finally:
        await upload.close()


@pytest.mark.asyncio
async def test_upload_route_rejects_busy_session_without_parsing_upload(monkeypatch):
    upload = _upload("rows.csv")
    close_state = _track_close(upload)
    request, request_state = _request(upload)

    async def fake_check_session_access(*_args, **_kwargs):
        return SimpleNamespace(
            is_active=True,
            is_processing=True,
            session=SimpleNamespace(pending_approval=None),
            hf_username="alice",
        )

    monkeypatch.setattr(agent, "_check_session_access", fake_check_session_access)

    with pytest.raises(HTTPException) as exc_info:
        await agent.upload_session_dataset(
            "s1",
            request,
            {
                "user_id": "u1",
                "username": "alice",
                agent.INTERNAL_HF_TOKEN_KEY: "hf-token",
            },
        )

    assert exc_info.value.status_code == 409
    assert request_state["form_called"] is False
    assert close_state["closed"] is False
    await upload.close()


@pytest.mark.asyncio
async def test_upload_route_appends_context_note_and_persists(monkeypatch):
    upload = _upload("rows.jsonl", b'{"text":"hi"}\n')
    close_state = _track_close(upload)
    request, request_state = _request(upload)
    messages = []
    persisted = []
    agent_session = SimpleNamespace(
        is_active=True,
        is_processing=False,
        session=SimpleNamespace(
            pending_approval=None,
            context_manager=SimpleNamespace(add_message=messages.append),
        ),
        hf_username="alice",
    )
    uploaded = dataset_uploads.DatasetUpload(
        session_id="s1",
        repo_id="alice/ml-intern-s1-datasets",
        repo_type="dataset",
        private=True,
        upload_id="abc123",
        config_name="upload_abc123",
        filename="rows.jsonl",
        original_filename="rows.jsonl",
        path_in_repo="uploads/abc123/train.jsonl",
        raw_path_in_repo="uploads/abc123/raw/rows.jsonl",
        normalized_path_in_repo="uploads/abc123/train.jsonl",
        normalized_format="jsonl",
        normalized_row_count=1,
        source_format="jsonl",
        source="session-upload",
        uploaded_at="2026-05-30T00:00:00Z",
        supports_training=True,
        size_bytes=14,
        format="jsonl",
        hub_url="https://huggingface.co/datasets/alice/ml-intern-s1-datasets/blob/main/uploads/abc123/train.jsonl",
        load_dataset_snippet=(
            "from datasets import load_dataset\n\n"
            'dataset = load_dataset("alice/ml-intern-s1-datasets", '
            '"upload_abc123", split="train", token=True)'
        ),
    )

    async def fake_check_session_access(*_args, **_kwargs):
        return agent_session

    async def fake_push_dataset_upload_to_hub(**kwargs):
        assert kwargs["upload"] is upload
        assert kwargs["hf_token"] == "hf-token"
        return uploaded

    async def fake_persist_session_snapshot(value):
        persisted.append(value)

    monkeypatch.setattr(agent, "_check_session_access", fake_check_session_access)
    monkeypatch.setattr(
        agent, "push_dataset_upload_to_hub", fake_push_dataset_upload_to_hub
    )
    monkeypatch.setattr(
        agent.session_manager,
        "persist_session_snapshot",
        fake_persist_session_snapshot,
    )

    response = await agent.upload_session_dataset(
        "s1",
        request,
        {
            "user_id": "u1",
            "username": "alice",
            agent.INTERNAL_HF_TOKEN_KEY: "hf-token",
        },
    )

    assert response.repo_id == uploaded.repo_id
    assert response.config_name == uploaded.config_name
    assert response.path_in_repo == uploaded.path_in_repo
    assert response.raw_path_in_repo == uploaded.raw_path_in_repo
    assert response.normalized_path_in_repo == uploaded.normalized_path_in_repo
    assert response.normalized_row_count == 1
    assert response.source_format == "jsonl"
    assert len(messages) == 1
    assert messages[0].role == "user"
    assert messages[0].content.startswith("[SYSTEM:")
    assert uploaded.config_name in messages[0].content
    assert "normalized dataset config" in messages[0].content
    assert uploaded.normalized_path_in_repo in messages[0].content
    assert uploaded.raw_path_in_repo in messages[0].content
    assert uploaded.load_dataset_snippet in messages[0].content
    assert "HF Jobs or GCP Vertex" in messages[0].content
    assert persisted == [agent_session]
    assert request_state["form_called"] is True
    assert close_state["closed"] is True


@pytest.mark.asyncio
async def test_upload_route_records_uploaded_dataset_metadata(monkeypatch):
    upload = _upload("notes.md", b"# Notes\n\nFine-tune on this.")
    request, _request_state = _request(upload)
    agent_session = SimpleNamespace(
        is_active=True,
        is_processing=False,
        session=SimpleNamespace(
            pending_approval=None,
            context_manager=SimpleNamespace(add_message=lambda _message: None),
            uploaded_datasets=[],
        ),
        hf_username="alice",
    )
    uploaded = dataset_uploads.DatasetUpload(
        session_id="s1",
        repo_id="alice/ml-intern-s1-datasets",
        repo_type="dataset",
        private=True,
        upload_id="md123",
        config_name="upload_md123",
        filename="notes.md",
        original_filename="notes.md",
        path_in_repo="uploads/md123/train.jsonl",
        raw_path_in_repo="uploads/md123/raw/notes.md",
        normalized_path_in_repo="uploads/md123/train.jsonl",
        normalized_format="jsonl",
        normalized_row_count=1,
        source_format="md",
        source="session-upload",
        uploaded_at="2026-05-30T00:00:00Z",
        supports_training=True,
        size_bytes=27,
        format="md",
        hub_url="https://huggingface.co/datasets/alice/ml-intern-s1-datasets/blob/main/uploads/md123/train.jsonl",
        load_dataset_snippet=(
            "from datasets import load_dataset\n\n"
            'dataset = load_dataset("alice/ml-intern-s1-datasets", '
            '"upload_md123", split="train", token=True)'
        ),
    )

    async def fake_check_session_access(*_args, **_kwargs):
        return agent_session

    async def fake_push_dataset_upload_to_hub(**_kwargs):
        return uploaded

    async def fake_persist_session_snapshot(_value):
        return None

    monkeypatch.setattr(agent, "_check_session_access", fake_check_session_access)
    monkeypatch.setattr(
        agent, "push_dataset_upload_to_hub", fake_push_dataset_upload_to_hub
    )
    monkeypatch.setattr(
        agent.session_manager,
        "persist_session_snapshot",
        fake_persist_session_snapshot,
    )

    try:
        response = await agent.upload_session_dataset(
            "s1",
            request,
            {
                "user_id": "u1",
                "username": "alice",
                agent.INTERNAL_HF_TOKEN_KEY: "hf-token",
            },
        )
    finally:
        await upload.close()

    assert response.source_format == "md"
    assert response.format == "md"
    assert agent_session.session.uploaded_datasets == [
        {
            "upload_id": "md123",
            "filename": "notes.md",
            "format": "md",
            "source_format": "md",
            "source": "session-upload",
            "uploaded_at": response.uploaded_at,
            "normalized_row_count": 1,
            "normalized_format": "jsonl",
            "status": "ready",
            "supports_training": True,
            "size_bytes": 27,
            "config_name": "upload_md123",
            "repo_id": "alice/ml-intern-s1-datasets",
            "repo_type": "dataset",
            "normalized_path_in_repo": "uploads/md123/train.jsonl",
            "raw_path_in_repo": "uploads/md123/raw/notes.md",
            "hub_url": uploaded.hub_url,
            "load_dataset_snippet": uploaded.load_dataset_snippet,
        }
    ]
    assert response.uploaded_at


@pytest.mark.asyncio
async def test_upload_route_closes_upload_when_hub_upload_fails(monkeypatch):
    upload = _upload("rows.csv")
    close_state = _track_close(upload)
    request, request_state = _request(upload)

    async def fake_check_session_access(*_args, **_kwargs):
        return SimpleNamespace(
            is_active=True,
            is_processing=False,
            session=SimpleNamespace(pending_approval=None),
            hf_username="alice",
        )

    async def fake_push_dataset_upload_to_hub(**_kwargs):
        raise RuntimeError("hub unavailable")

    monkeypatch.setattr(agent, "_check_session_access", fake_check_session_access)
    monkeypatch.setattr(
        agent, "push_dataset_upload_to_hub", fake_push_dataset_upload_to_hub
    )

    with pytest.raises(HTTPException) as exc_info:
        await agent.upload_session_dataset(
            "s1",
            request,
            {
                "user_id": "u1",
                "username": "alice",
                agent.INTERNAL_HF_TOKEN_KEY: "hf-token",
            },
        )

    assert exc_info.value.status_code == 502
    assert exc_info.value.detail == "Dataset upload failed. Please try again."
    assert request_state["form_called"] is True
    assert close_state["closed"] is True


@pytest.mark.asyncio
async def test_upload_route_maps_hub_permission_error_safely(monkeypatch):
    upload = _upload("rows.csv")
    close_state = _track_close(upload)
    request, request_state = _request(upload)

    async def fake_check_session_access(*_args, **_kwargs):
        return SimpleNamespace(
            is_active=True,
            is_processing=False,
            session=SimpleNamespace(pending_approval=None),
            hf_username="alice",
        )

    async def fake_push_dataset_upload_to_hub(**_kwargs):
        response = httpx.Response(
            403,
            request=httpx.Request("POST", "https://huggingface.co/api/datasets"),
            headers={"x-request-id": "req-123"},
        )
        raise HfHubHTTPError(
            "403 Forbidden: token hf_secret cannot write",
            response=response,
            server_message="token hf_secret cannot write",
        )

    monkeypatch.setattr(agent, "_check_session_access", fake_check_session_access)
    monkeypatch.setattr(
        agent, "push_dataset_upload_to_hub", fake_push_dataset_upload_to_hub
    )

    with pytest.raises(HTTPException) as exc_info:
        await agent.upload_session_dataset(
            "s1",
            request,
            {
                "user_id": "u1",
                "username": "alice",
                agent.INTERNAL_HF_TOKEN_KEY: "hf-token",
            },
        )

    assert exc_info.value.status_code == 403
    assert exc_info.value.detail == (
        "Hugging Face denied permission to create or write to the dataset repo."
    )
    assert "hf_secret" not in exc_info.value.detail
    assert request_state["form_called"] is True
    assert close_state["closed"] is True
